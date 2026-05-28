import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

# ═══════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数智光衡')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('高校教室自然光协同智能照明节能控制系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('可视化路演答辩稿 · 60 秒精讲版')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x00, 0x88, 0xCC)

doc.add_paragraph()

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# COMPETITION INFO
# ═══════════════════════════════════════════════════

info_text = """赛事：中国大学生计算机设计大赛 · 国赛（人工智能赛道）
作品类别：人工智能应用 — 智慧教育 + 碳中和
核心技术：YOLOv8n 视觉检测 + RK3588 NPU 边缘推理 + Open-Meteo 实时气象融合 + ECharts 多维数据可视化
演示地址：https://comethrue.github.io/luminous-carbon-atlas/
后端地址：https://luminous-carbon-atlas.onrender.com"""

p = doc.add_paragraph()
run = p.add_run(info_text)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x77, 0x88)
run.font.name = 'JetBrains Mono'

doc.add_paragraph()

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

# ═══════════════════════════════════════════════════
# PREPARATION
# ═══════════════════════════════════════════════════

h = doc.add_heading('录制前准备', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

prep_items = [
    'Chrome / Edge 无痕窗口打开演示地址，F12 → Network 确认 onrender.com 请求全部返回 200（绿色）后关闭 DevTools',
    '分辨率 1920×1080，浏览器缩放 100%，F11 全屏模式',
    'Windows：任务栏自动隐藏 | macOS：自动隐藏 Dock 和菜单栏',
    '光标移至屏幕右下角边缘（避免入镜），或使用光标隐藏工具',
    'OBS 录制参数：1080p / 30fps / 码率 8000kbps，关闭麦克风降噪，音量 -6dB 避削波',
    '建议录制 3 遍，选取滚动最流畅、节奏最好的一条',
]

for item in prep_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# SCRIPT TABLE
# ═══════════════════════════════════════════════════

h = doc.add_heading('答辩稿正文（精确到秒）', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

p = doc.add_paragraph()
run = p.add_run('标注：🖱 鼠标操作  ⏱ 计时节点  🎯 评审看点')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

doc.add_paragraph()

# ── SEGMENT 1 ──
h3 = doc.add_heading('第一段  0:00 – 0:10  开场定调', level=3)
for r in h3.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

operations = [
    ('🖱 画面', '页面完整加载后停在 Hero 首屏，全程不操作鼠标'),
    ('🎯 评审看点', '能量核心三层旋转光环 + 粒子背景 + 扫描线覆盖——三重视觉层次叠加，营造"碳观测舱"的沉浸感'),
]
for label, text in operations:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{text}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '各位评委老师好。我们带来的作品是"数智光衡"——基于 AI 视觉与自然光协同的 '
    '高校教室智能照明节能控制系统。\n\n'
    '全国高校教室照明年耗电量超过 400 亿千瓦时，其中约 40% 因"无人亮灯"和'
    '"自然光充足时仍全功率运行"而被白白浪费。传统照明系统是一个开环——它不'
    '知道教室里有没有人，也不知道窗外阳光正好。\n\n'
    '我们做的，就是给照明系统装上"眼睛"和"大脑"。'
)
run.font.size = Pt(10.5)

doc.add_paragraph()

# ── SEGMENT 2 ──
h3 = doc.add_heading('第二段  0:10 – 0:25  核心可视化 · 数据中枢', level=3)
for r in h3.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

operations = [
    ('🖱 0:10', '缓慢匀速向下滚动，依次露出 5 张 KPI 卡片（累计节电 → 节能率 → CO₂减排 → 等效植树 → 节省电费）'),
    ('🖱 0:14', '继续滚动至环境卡片 + 教室分区灯控图，停留 3 秒'),
    ('🖱 0:18', '滚动至 24h 功率对比曲线图，鼠标悬停在蓝色曲线"AI 控制"上方区域'),
    ('🎯 评审看点', '① 5 组 KPI 数字均有滚动动画（AnimatedNumber），从 0 递增至目标值 / ② 环境卡片右上角"Open-Meteo LIVE"红色脉冲指示灯，证明数据来自真实 API 非模拟 / ③ 教室俯视图三区独立控灯，红/绿/灰三色交替，对应 AI 分区决策逻辑'),
]
for label, text in operations:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{text}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '系统的核心架构分三层：感知层 —— USB 摄像头运行 YOLOv8n 实时检测人数与位置，'
    '同时接入 Open-Meteo 卫星气象数据获取太阳辐射和云量，全部免费、零密钥；'
    '决策层 —— Radxa Q6A 边缘计算板搭载 RK3588 NPU，6 TOPS 算力完成端侧推理，'
    '数据不出教室；执行层 —— 三路继电器独立控制左、中、右三个照明分区。\n\n'
    '现在看到的可视化大屏就是系统的数字孪生——每一组 KPI、每一段功率曲线、'
    '每一次分区切换，都是真实物理世界在数据空间的同步映射。\n\n'
    '注意右上角环境面板——"Open-Meteo LIVE"表示当前太阳辐射、云量、AQI 全部来自'
    '卫星实时数据。自然光充足时，系统自动降低照明功率；云层遮挡时即时补偿。'
    '这是系统"看天吃饭"的核心逻辑。'
)
run.font.size = Pt(10.5)

doc.add_paragraph()

# ── SEGMENT 3 ──
h3 = doc.add_heading('第三段  0:25 – 0:38  图表矩阵 · 多维论证', level=3)
for r in h3.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

operations = [
    ('🖱 0:25', '继续向下滚动，依次展示策略饼图 → 七日柱状图 → 场景雷达对比 → 校园能耗热力图'),
    ('🖱 0:32', '在场景对比图停留 3 秒，光标依次指向"全亮基线"（红色）→"AI 智能"（青色）'),
    ('🎯 评审看点', '① 24h 曲线对比：灰色全亮基线 vs 蓝色 AI 控制，二者间隙即为节能量——视觉上"一图胜千言" / ② 场景对比：四种方案并排，AI 方案年耗电最低 / ③ 校园热力图：中国地图标注 GDP 前十城市 + 广东试点，粒子流动动画'),
]
for label, text in operations:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{text}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '从单点 KPI 到多维交叉验证，我们用五类图表构建完整的数据证据链。\n\n'
    '24 小时功率曲线是核心——灰色区域代表传统全亮基线，恒定 480 瓦；蓝色曲线是'
    'AI 控制下的实际功率，与课表节奏和实时人数精确同步。上午 8 到 10 点高峰满功率，'
    '午休自动归零，下午随自然光增强逐步降低，精度达到 15 分钟级别。'
    '两条曲线之间的间隙面积，就是累计节电量。\n\n'
    '策略分布饼图展示了 AI 的决策逻辑——ALL_OFF 自动关灯、ZONE_PARTIAL 分区微亮、'
    'ALL_ON 全域点亮，三种策略的占比精确对应了一天的课表结构。\n\n'
    '场景对比直接给出结论：相比全亮基线，AI 方案年节电率 72%，年节省电费超过 '
    '八百元每间教室。'
)
run.font.size = Pt(10.5)

doc.add_paragraph()

# ── SEGMENT 4 ──
h3 = doc.add_heading('第四段  0:38 – 0:50  碳账本 · 从节能到碳中和', level=3)
for r in h3.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

operations = [
    ('🖱 0:38', '加速滚动掠过叙事章节（Chapter 01-05），不做停留'),
    ('🖱 0:42', '滚动至碳账本六宫格卡片，稍慢滚动逐张展示'),
    ('🖱 0:46', '滚动至中国地图全貌，停留至段尾'),
    ('🎯 评审看点', '① 每张碳账本卡片下方标注了国标/部委引用来源——数据有据可查 / ② 中国地图使用本地 GeoJSON 渲染（非在线瓦片），无网络依赖 / ③ 地图周围 Canvas 粒子轨道动画，营造"碳流"视觉效果'),
]
for label, text in operations:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{text}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '节能数据最终需要翻译成碳语言——因为比赛的终极命题不是"省了多少钱"，'
    '而是"减了多少碳"。\n\n'
    '碳账本模块将累计节电量换算为三重生态指标：CO₂ 减排量、标准煤节约量、等效植树量。'
    '每一项数据都引用了国家权威标准——碳排放因子来自生态环境部 2025 年公告，'
    '碳汇换算依据国家林草局 LY/T 2988 标准。每一度电的生态账，都有据可查。\n\n'
    '中国地图展示的是推广潜力：若全校 200 间教室全部接入系统，'
    '年减排二氧化碳超过 100 吨，等效植树约 4,600 棵。一间教室的优化，'
    '可以放大为整个校园的碳中和基础设施。'
)
run.font.size = Pt(10.5)

doc.add_paragraph()

# ── SEGMENT 5 ──
h3 = doc.add_heading('第五段  0:50 – 1:00  收尾升华', level=3)
for r in h3.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

operations = [
    ('🖱 0:50', '快速滚回页面顶部 Hero 区域，让"数智光衡"标题重新居中'),
    ('🖱 0:55', '页面静止不动，保持 Hero 画面至结束'),
]
for label, text in operations:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{text}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x80)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '总结三点创新：第一，我们首次将 AI 视觉检测与卫星气象数据在照明控制场景中'
    '做了深度融合，让系统真正实现"看天吃饭"；第二，端侧 NPU 推理实现了毫秒级响应'
    '与数据隐私保护，所有计算在教室本地完成；第三，我们构建了一套完整的数据可视化'
    '叙事体系，让抽象的算法决策变得直观、透明、有说服力。\n\n'
    '"数智光衡"——用 AI 和自然光，照亮每一间低碳教室。\n\n'
    '谢谢各位评委老师，请批评指正。'
)
run.font.size = Pt(10.5)

doc.add_paragraph()

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# COMPETITION STRATEGY
# ═══════════════════════════════════════════════════

h = doc.add_heading('国赛答辩策略要点', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

strategies = [
    ('差异化定位',
     '多数 AI 赛道作品侧重模型精度或算法创新，我们选择"可视化叙事"作为差异化切入点——'
     '将 AI 系统的工作过程用六类动态图表直观呈现，让评委在 60 秒内完整理解技术逻辑。'),
    ('技术可信度',
     '遇到"数据来源"问题时，直接打开 Network 面板证明数据来自 Open-Meteo 实时 API '
     '而非本地 Mock。碳账本每一项换算均标注国标编号，评委质疑时可直接引用。'),
    ('工程完整度',
     '作品涵盖端侧推理（RK3588 NPU）+ 后端服务（FastAPI）+ 前端可视化（React/ECharts）'
     '+ CI/CD 部署（GitHub Actions + Render），是一套完整的生产级系统而非原型。'),
    ('社会价值锚定',
     '紧扣"教育数字化"与"碳中和"两大国家战略，强调全校推广潜力和碳减排的规模化效益。'
     '评委中有教育信息化和绿色校园方向的专家，这两个关键词必须出现。'),
    ('常见追问准备',
     '① Q: 数据是否真实？→ A: 天气数据来自 Open-Meteo 卫星 API，碳排放因子引用生态环境部 2025 年公告，'
     '教室功率模型基于《高校教室使用率调研报告》\n'
     '② Q: 算法创新在哪？→ A: 自然光因子动态权重调整——将实时太阳辐射和云量输入控制策略，'
     '在人员检测基础上叠加光照约束，实现"有人+暗=亮，有人+亮=暗"的双条件决策\n'
     '③ Q: 硬件成本？→ A: Radxa Q6A 约 ¥600，USB 摄像头 ¥50，继电器模块 ¥30，'
     '单教室硬件成本 < ¥700，适合大规模推广\n'
     '④ Q: 与竞品差异？→ A: 传统传感器方案只能检测"有没有人"，无法判断"具体位置"和"自然光是否足够"；'
     '我们的 AI 视觉方案同时解决两个维度的问题'),
]

for title, text in strategies:
    p = doc.add_paragraph()
    run = p.add_run(f'● {title}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# CHECKLIST
# ═══════════════════════════════════════════════════

h = doc.add_heading('录制质量检查清单', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x0D, 0x16, 0x28)

checks = [
    '画面 1080p 无黑边、无缩放瑕疵、滚动无画面撕裂',
    'Network 面板 onrender.com 请求全部 200（录制前确认）',
    '环境卡片显示"Open-Meteo LIVE"带红色脉冲指示点',
    '顶部状态栏滚动播报条正常轮播',
    '所有 KPI 数字、碳账本数字均有递增动画效果',
    '中国地图 GeoJSON 正常加载，无灰色占位区域',
    'Canvas 粒子动画流畅（背景 + 地图周围轨道粒子）',
    'CSS 扫描线覆盖层持续运行',
    '全屏无任务栏、无 Dock、无浏览器标签栏、无鼠标光标',
    '音频清晰无喷麦、无环境噪音、语速稳定',
    '总时长 60±3 秒（建议目标 58 秒留余量）',
]

for item in checks:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PRESENTATION_SCRIPT.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
