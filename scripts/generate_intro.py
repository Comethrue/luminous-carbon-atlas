from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections: s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
style = doc.styles['Normal']; style.font.size=Pt(10.5); style.font.name='等线'

# ── Title ──
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('数智光衡 · 可视化大屏介绍稿件（精简版）'); r.bold=True; r.font.size=Pt(20)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('全程约 55 秒 ｜ 匀速向下滚动'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
doc.add_paragraph()

# ── S1 ──
h=doc.add_heading('1. Hero', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 页面停在顶部'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 数智光衡——AI视觉结合卫星气象数据，按三个分区独立控灯，实现教室照明"有人才亮、有光就暗"'); r.font.size=Pt(10.5)

# ── S2 ──
h=doc.add_heading('2. KPI + 自然光贡献', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 缓慢下滚，5张KPI卡 → 环境卡停2秒'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 累计节电、节能率、CO₂减排、等效植树、电费节省——五组指标实时计算。右侧自然光贡献率接入Open-Meteo卫星数据，太阳辐射和云量直接驱动AI控灯策略'); r.font.size=Pt(10.5)

# ── S3 ──
h=doc.add_heading('3. 教室灯控 + 能耗对比', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 继续下滚 → 三区灯控 → 24h曲线停3秒'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 教室俯视图实时定位人员位置，左中右三区独立亮灭。下方24小时功率曲线，灰色全亮基线恒定480瓦，蓝色AI控制按课表节奏波动——午休自动关闭，自然光充足时自动调暗，绿色区域即节电量'); r.font.size=Pt(10.5)

# ── S4 ──
h=doc.add_heading('4. 策略 + 场景对比', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 快速滚动，场景对比停2秒'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 全亮基线年耗1881度，AI方案仅527度——节能率72%'); r.font.size=Pt(10.5)

# ── S5 ──
h=doc.add_heading('5. 叙事章节（快速过）', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 加速滚动，不停留'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 五章叙事串联从浪费发现到碳账本换算的完整故事线'); r.font.size=Pt(10.5)

# ── S6 ──
h=doc.add_heading('6. 碳账本 + 中国地图', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 缓慢滚动，卡片逐一展示 → 地图停4秒'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 节电量换算为六项生态指标，每项标注国标出处。中国地图展示十城推广潜力，全校200间教室接入后年减排CO₂超100吨'); r.font.size=Pt(10.5)

# ── S7 ──
h=doc.add_heading('7. 收尾', level=2)
p=doc.add_paragraph(); r=p.add_run('🖱 滚回顶部'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x99,0xBB)
p=doc.add_paragraph(); r=p.add_run('🔊 数智光衡——用AI和自然光，照亮每一间低碳教室'); r.font.size=Pt(10.5)

doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('━'*40); r.font.size=Pt(8); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
doc.add_paragraph()

# ── Cheat card ──
h=doc.add_heading('📋 极简提词卡', level=2)
cheats = [
    'Hero        → AI视觉+卫星气象，三区独立控灯',
    'KPI+环境    → 五组指标实时计算，卫星数据驱动策略',
    '灯控+曲线   → 三区独立亮灭，AI vs 全亮基线，绿区即节电',
    '场景对比    → 全亮1881度 → AI仅527度，节能72%',
    '碳账本+地图 → 六项生态指标国标可查，全校年减排超100吨',
    '收尾        → 用AI和自然光照亮低碳教室',
]
for c in cheats:
    p=doc.add_paragraph(c); r=p.runs[0] if p.runs else None
    for r in p.runs: r.font.size=Pt(10); r.font.name='JetBrains Mono'

doc.add_paragraph()
h=doc.add_heading('✅ 录制前检查', level=2)
for i in ['Chrome无痕 → F12确认onrender.com请求200 → 关DevTools → F11全屏',
           '确认环境卡片显示"Open-Meteo LIVE"红色指示灯',
           '全程匀速滚动，不做回滚，总长控制在60秒内']:
    doc.add_paragraph(i, style='List Bullet')

output = 'C:/Users/a/Desktop/PCB/luminous-carbon-atlas/WEB_INTRO_V2.docx'
doc.save(output)
print(f'Done: {output}')
